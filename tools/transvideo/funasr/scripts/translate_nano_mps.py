#!/usr/bin/env python3
"""Transcribe video/audio with Fun-ASR-Nano-2512 and mandatory Apple MPS."""

from __future__ import annotations

import argparse
import hashlib
import importlib.util
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from datetime import datetime
from pathlib import Path
from typing import Any


DEFAULT_MODEL = "FunAudioLLM/Fun-ASR-Nano-2512"
DEFAULT_VAD_MODEL = "fsmn-vad"
DEFAULT_SPK_MODEL = "cam++"
DEFAULT_LANGUAGE = "中文"
MODEL_ROOT = Path("models")
MODELSCOPE_CACHE_ROOT = Path.home() / ".cache" / "modelscope" / "hub" / "models" / "iic"
MODELSCOPE_ALIASES = {
    "fsmn-vad": MODELSCOPE_CACHE_ROOT / "speech_fsmn_vad_zh-cn-16k-common-pytorch",
    "cam++": MODELSCOPE_CACHE_ROOT / "speech_campplus_sv_zh-cn_16k-common",
}
REMOTE_CODE = Path(__file__).resolve().parents[1] / "third_party" / "Fun-ASR" / "model.py"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Use FunASR Nano 2512 to turn local video/audio into text and subtitles. "
            "This script requires PyTorch MPS and will not fall back to CPU."
        )
    )
    parser.add_argument("input", type=Path, help="Video or audio file to process.")
    parser.add_argument(
        "-o",
        "--output-dir",
        type=Path,
        default=Path("output"),
        help="Root directory for txt/json/docx outputs. Default: output",
    )
    parser.add_argument(
        "--model",
        default=DEFAULT_MODEL,
        help="Hugging Face model id or local model path. Default: FunAudioLLM/Fun-ASR-Nano-2512",
    )
    parser.add_argument(
        "--hf-endpoint",
        default=os.environ.get("HF_ENDPOINT"),
        help="Optional Hugging Face endpoint/mirror, e.g. https://hf-mirror.com.",
    )
    parser.add_argument(
        "--vad-model",
        default=DEFAULT_VAD_MODEL,
        help="VAD model id/local path. Use 'none' to disable. Default: fsmn-vad",
    )
    parser.add_argument(
        "--spk-model",
        default=DEFAULT_SPK_MODEL,
        help="Speaker diarization model id/local path. Use 'none' to disable. Default: cam++",
    )
    parser.add_argument(
        "--language",
        default=DEFAULT_LANGUAGE,
        help="Recognition language for Nano 2512, e.g. 中文, 英文, 日文, 韩文, 粤语. Default: 中文",
    )
    parser.add_argument(
        "--batch-size-s",
        type=int,
        default=0,
        help=(
            "FunASR VAD batching window in seconds. Default: 0, which forces one VAD "
            "segment at a time because Fun-ASR-Nano-2512 does not implement batch decoding."
        ),
    )
    parser.add_argument(
        "--max-length",
        type=int,
        default=4096,
        help="Maximum new tokens generated for each VAD segment. Default: 4096",
    )
    parser.add_argument(
        "--speaker-threshold",
        type=float,
        default=0.4,
        help="Cosine threshold for custom speaker clustering. Lower merges more segments. Default: 0.4",
    )
    return parser.parse_args()


def require_ffmpeg() -> None:
    if shutil.which("ffmpeg") is None:
        raise RuntimeError("ffmpeg was not found. Install it first, for example: brew install ffmpeg")


def require_mps() -> None:
    try:
        import torch
    except Exception as exc:
        raise RuntimeError("PyTorch is required before MPS can be used.") from exc

    if not torch.backends.mps.is_available():
        raise RuntimeError(
            "PyTorch MPS is not available. This script is MPS-only and will not fall back to CPU."
        )


def require_nano_dependencies() -> None:
    missing = [
        package
        for package in ("transformers", "safetensors", "whisper", "tiktoken")
        if importlib.util.find_spec(package) is None
    ]
    if missing:
        raise RuntimeError(
            "Fun-ASR-Nano-2512 requires missing package(s): "
            f"{', '.join(missing)}. Install them with: pip install {' '.join(missing)}"
        )


def register_nano_remote_code(remote_code: Path = REMOTE_CODE) -> str:
    if not remote_code.exists():
        raise RuntimeError(
            f"Fun-ASR-Nano remote code was not found: {remote_code}. "
            "Run: git clone --depth 1 https://github.com/FunAudioLLM/Fun-ASR.git /tmp/Fun-ASR"
        )

    code_dir = str(remote_code.parent)
    if code_dir not in sys.path:
        sys.path.insert(0, code_dir)

    from funasr.utils.dynamic_import import import_module_from_path

    import_module_from_path(str(remote_code))
    return str(remote_code)


def file_md5(path: Path, chunk_size: int = 1024 * 1024) -> str:
    digest = hashlib.md5()
    with path.open("rb") as file_obj:
        for chunk in iter(lambda: file_obj.read(chunk_size), b""):
            digest.update(chunk)
    return digest.hexdigest()


def extract_audio(input_path: Path, wav_path: Path) -> None:
    cmd = [
        "ffmpeg",
        "-y",
        "-i",
        str(input_path),
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
        "-f",
        "wav",
        str(wav_path),
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE, text=True)


def prepare_wav(input_path: Path, dated_output_dir: Path) -> tuple[Path, str, str, str]:
    source_md5 = file_md5(input_path)
    if input_path.suffix.lower() == ".wav":
        wav_path = input_path
        wav_md5 = file_md5(wav_path)
    else:
        temp_wav = dated_output_dir / f"{input_path.stem}.{source_md5}.wav"
        print(f"[1/4] Converting to 16 kHz mono WAV: {temp_wav}")
        extract_audio(input_path, temp_wav)
        wav_md5 = file_md5(temp_wav)
        wav_path = dated_output_dir / f"{input_path.stem}.{source_md5}.{wav_md5}.wav"
        if wav_path != temp_wav:
            temp_wav.replace(wav_path)

    output_stem = f"{input_path.stem}.{source_md5}.{wav_md5}"
    return wav_path, output_stem, source_md5, wav_md5


def normalize_model(value: str) -> str | None:
    return None if value.lower() in {"", "none", "null", "false"} else value


def prefer_local_model(value: str | None, model_root: Path = MODEL_ROOT) -> str | None:
    if value is None:
        return None

    configured_path = Path(value).expanduser()
    if configured_path.exists():
        return str(configured_path.resolve())

    local_snapshot = model_root / value
    if local_snapshot.exists():
        return str(local_snapshot.resolve())

    modelscope_snapshot = MODELSCOPE_ALIASES.get(value)
    if modelscope_snapshot and modelscope_snapshot.exists():
        return str(modelscope_snapshot.resolve())

    return value


def collect_text(result: list[dict[str, Any]]) -> str:
    chunks: list[str] = []
    for item in result:
        text = item.get("text")
        if text:
            chunks.append(str(text).strip())
    return "\n".join(chunks).strip()


def clean_transcript_text(text: str) -> str:
    text = re.sub(r"<\|[^|]+?\|>", "", text)
    return re.sub(r"\s+", " ", text).strip()


def ms_to_clock_time(ms: int | float) -> str:
    total_ms = int(ms)
    hours, remainder = divmod(total_ms, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    seconds = remainder // 1_000
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def format_elapsed(seconds: float) -> str:
    whole_seconds = int(seconds)
    hours, remainder = divmod(whole_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{int((seconds - whole_seconds) * 100):02d}"


def extract_segments(result: list[dict[str, Any]]) -> list[dict[str, Any]]:
    segments: list[dict[str, Any]] = []
    speaker_roles: dict[str, str] = {}

    def role_for(raw_speaker: Any) -> str:
        speaker = str(raw_speaker if raw_speaker is not None else 0)
        if speaker not in speaker_roles:
            speaker_roles[speaker] = f"角色{len(speaker_roles) + 1}"
        return speaker_roles[speaker]

    for item in result:
        sentence_info = item.get("sentence_info") or []
        for sentence in sentence_info:
            text = str(sentence.get("text") or sentence.get("sentence") or "").strip()
            text = clean_transcript_text(text)
            if not text:
                continue
            start = int(sentence.get("start", 0))
            end = int(sentence.get("end", start + 1))
            role = role_for(sentence.get("spk", 0))
            segments.append({"start": start, "end": end, "role": role, "text": text})
        if sentence_info:
            continue

        text = clean_transcript_text(str(item.get("text", "")).strip())
        if not text:
            continue

        timestamps = item.get("timestamp") or []
        if timestamps:
            start = int(timestamps[0][0])
            end = int(timestamps[-1][1])
        else:
            start, end = 0, 1_000

        segments.append({"start": start, "end": end, "role": role_for(item.get("spk", 0)), "text": text})
    return segments


def format_segment(segment: dict[str, Any]) -> str:
    return (
        f"[{ms_to_clock_time(segment['start'])} – {ms_to_clock_time(segment['end'])}] "
        f"{segment['role']}：{segment['text']}"
    )


def write_docx(path: Path, segments: list[dict[str, Any]], title: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for segment in segments:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        time_run = paragraph.add_run(
            f"[{ms_to_clock_time(segment['start'])} – {ms_to_clock_time(segment['end'])}] "
        )
        time_run.bold = True
        role_run = paragraph.add_run(f"{segment['role']}：")
        role_run.bold = True
        paragraph.add_run(segment["text"])

    document.save(path)


def output_path(base: Path, extension: str) -> Path:
    return base.parent / f"{base.name}.{extension}"


def slice_audio(audio: Any, sample_rate: int, start_ms: int, end_ms: int) -> Any:
    start = max(int(start_ms * sample_rate / 1000), 0)
    end = min(int(end_ms * sample_rate / 1000), len(audio))
    return audio[start:end]


def assign_speaker_labels(embeddings: list[Any], threshold: float) -> list[int]:
    import numpy as np

    if not embeddings:
        return []

    vectors = np.stack([embedding.detach().cpu().numpy().reshape(-1) for embedding in embeddings])
    norms = np.linalg.norm(vectors, axis=1, keepdims=True)
    vectors = vectors / np.maximum(norms, 1e-12)

    centroids: list[Any] = []
    counts: list[int] = []
    labels: list[int] = []
    for vector in vectors:
        if not centroids:
            centroids.append(vector.copy())
            counts.append(1)
            labels.append(0)
            continue

        scores = [float(vector @ centroid) for centroid in centroids]
        best_label = int(np.argmax(scores))
        if scores[best_label] >= threshold:
            labels.append(best_label)
            counts[best_label] += 1
            centroids[best_label] = (
                centroids[best_label] * (counts[best_label] - 1) + vector
            ) / counts[best_label]
            centroids[best_label] = centroids[best_label] / max(
                np.linalg.norm(centroids[best_label]), 1e-12
            )
        else:
            labels.append(len(centroids))
            centroids.append(vector.copy())
            counts.append(1)

    return labels


def generate_diarized_result(
    *,
    wav_path: Path,
    asr_model: Any,
    vad_model: Any,
    spk_model: Any,
    language: str,
    max_length: int,
    speaker_threshold: float,
) -> list[dict[str, Any]]:
    import soundfile as sf

    vad_result = vad_model.generate(input=str(wav_path))
    vad_segments = (vad_result[0].get("value") if vad_result else []) or []
    if not vad_segments:
        return []

    audio, sample_rate = sf.read(str(wav_path), dtype="float32")
    if getattr(audio, "ndim", 1) > 1:
        audio = audio[:, 0]

    segments: list[dict[str, Any]] = []
    embeddings: list[Any] = []
    with tempfile.TemporaryDirectory(prefix="funasr_segments_") as temp_dir:
        temp_root = Path(temp_dir)
        total = len(vad_segments)
        for index, (start_ms, end_ms) in enumerate(vad_segments, start=1):
            segment_audio = slice_audio(audio, sample_rate, int(start_ms), int(end_ms))
            if len(segment_audio) == 0:
                continue

            segment_wav = temp_root / f"segment_{index:04d}.wav"
            sf.write(str(segment_wav), segment_audio, sample_rate)

            asr_result = asr_model.generate(
                input=[str(segment_wav)],
                language=language,
                batch_size=1,
                max_length=max_length,
            )
            text = clean_transcript_text(collect_text(asr_result))
            if not text:
                print(f"[3/4] Segment {index}/{total}: empty, skipped")
                continue

            spk_result = spk_model.generate(input=[segment_audio])
            embedding = spk_result[0].get("spk_embedding") if spk_result else None
            if embedding is None:
                raise RuntimeError(f"Speaker embedding was not returned for segment {index}.")

            embeddings.append(embedding)
            segments.append(
                {
                    "start": int(start_ms),
                    "end": int(end_ms),
                    "text": text,
                    "spk": 0,
                }
            )
            print(f"[3/4] Segment {index}/{total}: {ms_to_clock_time(start_ms)}-{ms_to_clock_time(end_ms)}")

    labels = assign_speaker_labels(embeddings, speaker_threshold)
    for segment, label in zip(segments, labels):
        segment["spk"] = int(label)

    return [
        {
            "key": wav_path.stem,
            "text": " ".join(segment["text"] for segment in segments),
            "sentence_info": [
                {
                    "start": segment["start"],
                    "end": segment["end"],
                    "sentence": segment["text"],
                    "spk": segment["spk"],
                }
                for segment in segments
            ],
            "diarization": {
                "method": "custom_vad_camplusplus_online_cosine",
                "speaker_threshold": speaker_threshold,
                "speaker_count": len(set(labels)) if labels else 0,
                "vad_segment_count": len(vad_segments),
                "transcribed_segment_count": len(segments),
            },
        }
    ]


def write_outputs(
    base: Path,
    result: list[dict[str, Any]],
    source_path: Path,
    source_md5: str,
    wav_md5: str,
) -> None:
    segments = extract_segments(result)
    if not segments:
        text = collect_text(result)
        if text:
            segments = [{"start": 0, "end": 1_000, "role": "角色1", "text": text}]

    transcript = "\n".join(format_segment(segment) for segment in segments)
    output_path(base, "txt").write_text(transcript + ("\n" if transcript else ""), encoding="utf-8")

    payload = {
        "source_file": str(source_path),
        "source_md5": source_md5,
        "wav_md5": wav_md5,
        "segments": segments,
        "raw_result": result,
    }
    output_path(base, "json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    write_docx(output_path(base, "docx"), segments, source_path.name)


def main() -> int:
    started_at = time.perf_counter()
    args = parse_args()
    if args.hf_endpoint:
        os.environ["HF_ENDPOINT"] = args.hf_endpoint

    input_path = args.input.expanduser().resolve()
    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 2

    try:
        require_ffmpeg()
        require_mps()
        require_nano_dependencies()
        remote_code = register_nano_remote_code()
    except RuntimeError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 2

    dated_output_dir = args.output_dir / datetime.now().strftime("%Y%m%d")
    dated_output_dir.mkdir(parents=True, exist_ok=True)

    wav_for_asr, output_stem, source_md5, wav_md5 = prepare_wav(input_path, dated_output_dir)
    output_base = dated_output_dir / output_stem

    print("[2/4] Loading Fun-ASR-Nano-2512 on mps. First run may download model files.")
    from funasr import AutoModel

    asr_model = prefer_local_model(args.model)
    vad_model = prefer_local_model(normalize_model(args.vad_model))
    spk_model = prefer_local_model(normalize_model(args.spk_model))

    model_kwargs: dict[str, Any] = {
        "model": asr_model,
        "hub": "hf",
        "device": "mps",
        "disable_update": True,
        "trust_remote_code": True,
        "remote_code": remote_code,
    }

    asr = AutoModel(**model_kwargs)

    if spk_model:
        if not vad_model:
            raise RuntimeError("Speaker diarization requires --vad-model; do not use --vad-model none.")
        print(f"[2/4] Loading VAD model on mps: {vad_model}")
        vad = AutoModel(model=vad_model, device="mps", disable_update=True)
        print(f"[2/4] Loading speaker model on mps: {spk_model}")
        spk = AutoModel(model=spk_model, device="mps", disable_update=True)
        print(
            f"[3/4] Recognizing and diarizing with language={args.language!r} "
            f"on mps, threshold={args.speaker_threshold}..."
        )
        result = generate_diarized_result(
            wav_path=wav_for_asr,
            asr_model=asr,
            vad_model=vad,
            spk_model=spk,
            language=args.language,
            max_length=args.max_length,
            speaker_threshold=args.speaker_threshold,
        )
    else:
        if vad_model:
            model_kwargs["vad_model"] = vad_model
            asr = AutoModel(**model_kwargs)
        print(f"[3/4] Recognizing with language={args.language!r} on mps...")
        result = asr.generate(
            input=str(wav_for_asr),
            language=args.language,
            batch_size=1,
            batch_size_s=args.batch_size_s,
            max_length=args.max_length,
        )
    print("[4/4] Writing txt/json/docx outputs...")
    write_outputs(output_base, result, input_path, source_md5, wav_md5)

    print(f"TXT : {output_path(output_base, 'txt')}")
    print(f"JSON: {output_path(output_base, 'json')}")
    print(f"DOCX: {output_path(output_base, 'docx')}")
    elapsed = time.perf_counter() - started_at
    print(f"耗时: {format_elapsed(elapsed)} ({elapsed:.2f}s)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
